import streamlit as st
import pandas as pd
import os 
from io import BytesIO
from docx import Document
from fpdf import FPDF

st.set_page_config(page_title="Data Sweeper", layout='wide')
st.title("Data Sweeper")
st.write("Transform your files between CSV, Excel, PDF, and Word formats effortlessly with built-in data cleaning and visualization!")

uploaded_file = st.file_uploader("Upload your CSV, Excel, PDF, or Word file", type=["csv", "xlsx", "xls", "pdf", "docx"],
accept_multiple_files=True)

if uploaded_file:
    for file in uploaded_file:
        file_ext = os.path.splitext(file.name)[-1].lower()
        file_name = file.name

        if file_ext == ".csv":
            df = pd.read_csv(file)
        elif file_ext == ".xlsx":
            df = pd.read_excel(file)
        elif file_ext == ".docx":
            st.warning("Word files can only be converted, not previewed.")
            df = pd.DataFrame()
        else:
            st.error(f"Unsupported file type: {file_ext}")
            continue

        if not df.empty:
            st.write(f"**File name:** {file.name}")
            st.write(f"**File Size:** {file.size / 1024:.2f} KB")
            st.write("Preview the head of the Dataframe")
            st.dataframe(df.head())

            st.subheader("Data Cleaning Options")
            if st.checkbox(f"Clean Data for {file.name}"):
                col1, col2 = st.columns(2)
                
                with col1:
                    if st.button(f"Remove Duplicates from {file.name}"):
                        df.drop_duplicates(inplace=True)
                        st.write("Duplicates Removed")
                
                with col2:
                    if st.button(f"Fill Missing Values for {file.name}"):
                        numeric_cols = df.select_dtypes(include=['number']).columns
                        df[numeric_cols] = df[numeric_cols].fillna(df[numeric_cols].mean())
                        st.write("Missing Values have been Filled!")

            st.subheader("Select Columns to Convert")
            columns = st.multiselect(f"Choose Columns for {file.name}", df.columns, default=df.columns)
            df = df[columns]

        st.subheader("Conversion Options")
        Conversion_type = st.radio(f"Convert {file.name} to:", ["CSV", "Excel", "Word (DOCX)", "PDF"])
        buffer = BytesIO()

        if not file_name:
            file_name = "converted_file"

        if Conversion_type == "CSV":
            df.to_csv(buffer, index=False)
            file_name = file_name.replace(file_ext, ".csv")
            mime_type = "text/csv"
        
        elif Conversion_type == "Excel":
            df.to_excel(buffer, index=False)
            file_name = file_name.replace(file_ext, ".xlsx")
            mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

        elif Conversion_type == "Word (DOCX)":
            doc = Document()
            doc.add_heading("Converted Data", level=1)
            
            if not df.empty:
                table = doc.add_table(rows=1, cols=len(df.columns))
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(df.columns):
                    hdr_cells[i].text = col_name
                
                for index, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
            
            doc.save(buffer)
            file_name = file_name.replace(file_ext, ".docx")
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        elif Conversion_type == "PDF":
            if file_ext == ".docx":
                doc = Document(file)
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("Arial", size=12)
                
                for para in doc.paragraphs:
                    pdf.multi_cell(0, 10, para.text)
                
                pdf_output = pdf.output(dest='S').encode('latin1')
                buffer.write(pdf_output)
                buffer.seek(0)
                file_name = file_name.replace(file_ext, ".pdf")
                mime_type = "application/pdf"
            else:
                st.warning("PDF conversion is only supported for Word files right now.")
                mime_type = "application/pdf"
        
        buffer.seek(0)

        st.download_button(
            label=f"Download {file.name} as {Conversion_type}",
            data=buffer,
            file_name=file_name,
            mime=mime_type,
        )

st.success("All files processed!")

st.markdown("---")
st.markdown("Created by Saboor Haider ❤️")
